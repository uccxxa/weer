from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
import os, re, uuid, time, traceback, json
import urllib.parse
from datetime import datetime
import requests
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage, Table, TableStyle
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.units import cm
from reportlab.lib.colors import black, HexColor
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = FastAPI(title="Mustaqil Ish Generator API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

GITHUB_TOKEN = os.getenv("GITHUB_TOKEN")
STABILITY_API_KEY = os.getenv("STABILITY_API_KEY")

if not GITHUB_TOKEN:
    print("WARNING: GITHUB_TOKEN topilmadi!")
if not STABILITY_API_KEY:
    print("WARNING: STABILITY_API_KEY topilmadi!")

TEMP_DIR = "/tmp/mustaqil_ish_files"
IMAGES_DIR = f"{TEMP_DIR}/images"

os.makedirs(TEMP_DIR, exist_ok=True)
os.makedirs(IMAGES_DIR, exist_ok=True)


# ─────────────────────────────────────────────────────
# MATNDAN MARKDOWN BELGILARINI TOZALASH
# ─────────────────────────────────────────────────────
def clean_text(text: str) -> str:
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
    text = re.sub(r'_{1,3}(.*?)_{1,3}', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)
    text = re.sub(r'`{1,3}(.*?)`{1,3}', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'^[-=]{3,}$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def clean_references(text: str) -> str:
    lines = text.split('\n')
    clean_lines = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        line = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', line)
        line = re.sub(r'_{1,3}(.*?)_{1,3}', r'\1', line)
        line = re.sub(r'^#{1,6}\s*', '', line)
        if re.match(r'^(REFERENCES|References|ADABIYOTLAR|Adabiyotlar):?$', line):
            continue
        clean_lines.append(line)
    return '\n'.join(clean_lines)


# ─────────────────────────────────────────────────────
# RASM GENERATSIYASI — Stability AI v2beta
# ─────────────────────────────────────────────────────
def generate_hf_image(prompt: str, filename: str) -> str | None:
    if not STABILITY_API_KEY:
        print("WARNING: STABILITY_API_KEY topilmadi!")
        return None
    try:
        print(f"Rasm generatsiyasi: {prompt[:60]}...")
        response = requests.post(
            "https://api.stability.ai/v2beta/stable-image/generate/core",
            headers={
                "Authorization": f"Bearer {STABILITY_API_KEY}",
                "Accept": "image/*"
            },
            files={"none": ""},
            data={
                "prompt": prompt[:500],
                "output_format": "png",
                "aspect_ratio": "16:9",
            },
            timeout=120
        )
        if response.status_code == 429:
            wait = int(response.headers.get("x-ratelimit-reset", 30))
            print(f"Rate limit. {wait}s kutish...")
            time.sleep(wait)
            response = requests.post(
                "https://api.stability.ai/v2beta/stable-image/generate/core",
                headers={
                    "Authorization": f"Bearer {STABILITY_API_KEY}",
                    "Accept": "image/*"
                },
                files={"none": ""},
                data={
                    "prompt": prompt[:500],
                    "output_format": "png",
                    "aspect_ratio": "16:9",
                },
                timeout=120
            )
        if response.status_code == 200:
            filepath = f"{IMAGES_DIR}/{filename}"
            with open(filepath, "wb") as f:
                f.write(response.content)
            print(f"Rasm saqlandi: {filepath}")
            return filepath
        else:
            print(f"Rasm xatosi ({response.status_code}): {response.text[:150]}")
            return None
    except Exception as e:
        print(f"Rasm xatosi: {e}")
        return None


def generate_topic_image(subject: str, topic: str) -> str | None:
    prompt = f"Professional academic illustration for {subject} topic: {topic}. Clean, scholarly, high-quality concept visualization. No text overlay."
    return generate_hf_image(prompt, "topic_main.png")


def generate_section_image(subject: str, section_title: str, idx: int) -> str | None:
    prompt = f"Academic infographic for {subject}: {section_title}. Professional diagram, scientific style, white background. No text overlay."
    return generate_hf_image(prompt, f"section_{idx}.png")


def generate_diagram_images(subject: str, topic: str) -> list:
    diagram_prompts = [
        (f"Flowchart diagram showing process flow of {topic} in {subject}. Clean lines, boxes, arrows, white background.", "diagram_1.png"),
        (f"Bar chart data visualization of {topic} in {subject}. Professional infographic, blue color scheme, white background.", "diagram_2.png"),
        (f"Hierarchical mind map of {topic} key concepts in {subject}. Academic style, clear structure, white background.", "diagram_3.png"),
        (f"Timeline diagram showing evolution of {topic} in {subject}. Horizontal flow, milestones, academic infographic.", "diagram_4.png"),
        (f"System architecture diagram of {topic} in {subject}. Technical blueprint style, clean lines.", "diagram_5.png"),
    ]
    paths = []
    for prompt, filename in diagram_prompts:
        paths.append(generate_hf_image(prompt, filename))
    return paths


def generate_extra_images(subject: str, topic: str) -> list:
    extra_prompts = [
        (f"Real-world application of {topic} in modern technology. Professional photograph, high quality.", "extra_1.png"),
        (f"Scientific laboratory research environment related to {topic}. Professional atmosphere, photorealistic.", "extra_2.png"),
        (f"Abstract conceptual digital art representing {topic} principles. Deep blue and gold color palette.", "extra_3.png"),
        (f"Futuristic visualization of innovations in {topic} field. Clean futuristic design.", "extra_4.png"),
    ]
    paths = []
    for prompt, filename in extra_prompts:
        paths.append(generate_hf_image(prompt, filename))
    return paths


# ─────────────────────────────────────────────────────
# MATN GENERATSIYASI
# ─────────────────────────────────────────────────────
def generate_text_with_retry(prompt: str, section_name: str, max_attempts: int = 3) -> str:
    for attempt in range(max_attempts):
        try:
            resp = requests.post(
                "https://models.inference.ai.azure.com/chat/completions",
                headers={"Authorization": f"Bearer {GITHUB_TOKEN}", "Content-Type": "application/json"},
                json={
                    "model": "gpt-4o-mini",
                    "messages": [
                        {"role": "system", "content": "You are an academic writer. Write plain text only. No markdown, no bold, no headers with # symbols, no bullet points with *, no special formatting. Use only plain paragraphs and numbered lists where needed."},
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": 0.4,
                    "max_tokens": 8000
                },
                timeout=300
            )
            if resp.status_code == 429:
                wait = int(resp.headers.get("x-ratelimit-reset", 20))
                print(f"Rate limit. {wait}s kutish...")
                time.sleep(wait)
                continue
            resp.raise_for_status()
            content = resp.json()["choices"][0]["message"]["content"].strip()
            content = clean_text(content)
            print(f"{section_name} tayyor: {len(content.split())} so'z")
            return content
        except Exception as e:
            print(f"{section_name} xatosi: {e}")
            if attempt == max_attempts - 1:
                return f"[{section_name} generatsiya qilinmadi]"
    return "[Xatolik]"


def generate_full_independent_work(subject: str, topic: str, cefr: str = "B2"):
    print(f"MUSTAQIL ISH: {subject} | {topic}")

    intro_prompt = f"""Write INTRODUCTION for independent work on "{topic}" in {subject}.
Language: English, CEFR {cefr} level.
Include: relevance, purpose, object and subject of research, research methods.
Requirements: 800-1000 words, formal academic style, plain text only, no markdown."""
    introduction = generate_text_with_retry(intro_prompt, "Kirish")

    outline_prompt = f"""Generate exactly 3 chapter titles for independent work on "{topic}" in {subject}.
Return ONLY a JSON array with 3 strings."""
    try:
        resp = requests.post(
            "https://models.inference.ai.azure.com/chat/completions",
            headers={"Authorization": f"Bearer {GITHUB_TOKEN}", "Content-Type": "application/json"},
            json={"model": "gpt-4o-mini", "messages": [{"role": "system", "content": "JSON only."}, {"role": "user", "content": outline_prompt}], "temperature": 0.3, "max_tokens": 300},
            timeout=60
        )
        raw = resp.json()["choices"][0]["message"]["content"]
        raw = re.sub(r'^```(?:json)?\n', '', raw, flags=re.IGNORECASE).strip()
        raw = re.sub(r'\n```$', '', raw).strip()
        chapter_titles = json.loads(raw)
        if len(chapter_titles) != 3:
            raise ValueError("3 ta sarlavha kerak")
    except Exception as e:
        print(f"Reja xatosi: {e}")
        chapter_titles = [
            f"Theoretical Foundations of {topic}",
            f"Analysis of {topic}",
            f"Practical Applications of {topic}"
        ]

    chapters = []
    for i, chapter_title in enumerate(chapter_titles, 1):
        print(f"Bob {i}: {chapter_title[:50]}...")
        chapter_prompt = f"""Write CHAPTER {i} for independent work.
Chapter Title: {chapter_title}
Main Topic: {topic}
Subject: {subject}
Language: English, CEFR {cefr} level.
Requirements:
- Minimum 2500 words
- Deep theoretical analysis with real examples
- Plain text only, no markdown, no # headers, no ** bold
- Use only numbered lists if needed"""
        chapter_content = generate_text_with_retry(chapter_prompt, f"Bob {i}")
        img_path = generate_section_image(subject, chapter_title, i)
        chapters.append({
            "title": chapter_title,
            "content": chapter_content,
            "words": len(chapter_content.split()),
            "image": img_path
        })

    conclusion_prompt = f"""Write CONCLUSION for independent work on "{topic}" in {subject}.
Include: summary of main findings, practical significance, recommendations.
Requirements: 600-800 words, formal academic style, plain text only, no markdown, no bold, no # symbols."""
    conclusion = generate_text_with_retry(conclusion_prompt, "Xulosa")

    references_prompt = f"""Generate REFERENCES for independent work on "{topic}" in {subject}.
Create exactly 15 academic references in APA style.
Return as plain numbered list only. Example:
1. Author, A. (2020). Title. Publisher.
No markdown, no bold, no extra text."""
    references_raw = generate_text_with_retry(references_prompt, "Adabiyotlar")
    references = clean_references(references_raw)

    main_image = generate_topic_image(subject, topic)
    diagrams = generate_diagram_images(subject, topic)
    extras = generate_extra_images(subject, topic)

    return {
        "introduction": introduction,
        "chapters": chapters,
        "conclusion": conclusion,
        "references": references,
        "main_image": main_image,
        "diagrams": diagrams,
        "extras": extras,
        "total_words": len(introduction.split()) + sum(c["words"] for c in chapters) + len(conclusion.split())
    }


# ─────────────────────────────────────────────────────
# PDF YARATISH
# ─────────────────────────────────────────────────────
def create_pdf(work_data: dict, meta: dict, pdf_path: str):
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        leftMargin=2*cm,
        rightMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()

    base_style = ParagraphStyle(
        'BaseStyle',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=14,
        leading=21,
        alignment=TA_JUSTIFY,
        spaceAfter=0,
        spaceBefore=0
    )
    title_style = ParagraphStyle(
        'Title',
        parent=base_style,
        fontName='Times-Bold',
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=0,
        spaceBefore=0
    )
    chapter_style = ParagraphStyle(
        'Chapter',
        parent=base_style,
        fontName='Times-Bold',
        fontSize=14,
        alignment=TA_LEFT,
        spaceAfter=0,
        spaceBefore=0
    )
    section_style = ParagraphStyle(
        'Section',
        parent=base_style,
        fontName='Times-Bold',
        fontSize=14,
        alignment=TA_CENTER,
        spaceAfter=0,
        spaceBefore=0
    )
    caption_style = ParagraphStyle(
        'Caption',
        parent=base_style,
        fontName='Times-Italic',
        fontSize=11,
        alignment=TA_CENTER,
        spaceAfter=0,
        spaceBefore=0,
        textColor=HexColor('#555555')
    )
    toc_style = ParagraphStyle(
        'TOCStyle',
        parent=base_style,
        fontName='Times-Roman',
        fontSize=14,
        leading=16,
        alignment=TA_LEFT,
        spaceAfter=0,
        spaceBefore=0
    )
    ref_style = ParagraphStyle(
        'RefStyle',
        parent=base_style,
        fontSize=12,
        leading=18,
        leftIndent=0.5*cm,
        spaceAfter=0,
        spaceBefore=0
    )

    elements = []
    LINE = Spacer(1, 14)

    # ═══ TITUL VARAQ ═══
    elements.append(Spacer(1, 2*cm))
    elements.append(Paragraph("MINISTRY OF HIGHER EDUCATION, SCIENCE AND INNOVATIONS",
                               ParagraphStyle('M1', parent=title_style, fontSize=12)))
    elements.append(LINE)
    elements.append(Paragraph("OF THE REPUBLIC OF UZBEKISTAN",
                               ParagraphStyle('M2', parent=title_style, fontSize=12)))
    elements.append(Spacer(1, 1.5*cm))
    elements.append(Paragraph(meta.get("university", "UNIVERSITY NAME").upper(),
                               ParagraphStyle('Uni', parent=title_style, fontSize=14)))
    elements.append(LINE)
    elements.append(Paragraph(meta.get("faculty", "FACULTY").upper(),
                               ParagraphStyle('Fac', parent=base_style, fontSize=12, alignment=TA_CENTER, spaceAfter=0)))
    elements.append(LINE)
    elements.append(Paragraph(f"Department: {meta.get('department', 'DEPARTMENT').upper()}",
                               ParagraphStyle('Dep', parent=base_style, fontSize=12, alignment=TA_CENTER, spaceAfter=0)))
    elements.append(Spacer(1, 2*cm))
    elements.append(Paragraph("INDEPENDENT WORK", title_style))
    elements.append(LINE)
    elements.append(Paragraph(f"Subject: {meta.get('subject', 'SUBJECT').upper()}",
                               ParagraphStyle('Subj', parent=base_style, fontSize=12, alignment=TA_CENTER, spaceAfter=0)))
    elements.append(LINE)
    elements.append(Paragraph(f"Topic: {meta.get('topic', 'TOPIC').upper()}",
                               ParagraphStyle('Topic', parent=base_style, fontName='Times-Bold', fontSize=13, alignment=TA_CENTER, spaceAfter=0)))
    elements.append(Spacer(1, 1.5*cm))

    if work_data.get('main_image') and os.path.exists(work_data['main_image']):
        try:
            img = RLImage(work_data['main_image'], width=14*cm, height=8*cm)
            img.hAlign = 'CENTER'
            elements.append(img)
            elements.append(LINE)
        except Exception as e:
            print(f"Asosiy rasm xatosi: {e}")

    elements.append(Paragraph(f"Performed by: {meta.get('student', 'STUDENT NAME').upper()}",
                               ParagraphStyle('Stud', parent=base_style, fontSize=12, alignment=TA_CENTER, spaceAfter=0)))
    elements.append(LINE)
    elements.append(Paragraph(f"Group: {meta.get('group', 'GROUP')}, Course: {meta.get('course', '1')}",
                               ParagraphStyle('Grp', parent=base_style, fontSize=12, alignment=TA_CENTER, spaceAfter=0)))
    elements.append(LINE)
    elements.append(Paragraph(f"Supervised by: {meta.get('teacher', 'TEACHER NAME').upper()}",
                               ParagraphStyle('Tch', parent=base_style, fontSize=12, alignment=TA_CENTER, spaceAfter=0)))
    elements.append(Spacer(1, 3*cm))
    elements.append(Paragraph(f"{meta.get('city', 'CITY').upper()} - {datetime.now().year}",
                               ParagraphStyle('City', parent=base_style, fontSize=12, alignment=TA_CENTER, spaceAfter=0)))
    elements.append(PageBreak())

    # ═══ MUNDARIJA ═══
    elements.append(Paragraph("TABLE OF CONTENTS", section_style))
    elements.append(LINE)

    toc_entries = [("1.", "INTRODUCTION")]
    for idx, ch in enumerate(work_data['chapters'], 2):
        toc_entries.append((f"{idx}.", ch['title'].upper()))
    toc_entries.append((f"{len(work_data['chapters'])+2}.", "CONCLUSION"))
    toc_entries.append((f"{len(work_data['chapters'])+3}.", "REFERENCES"))

    toc_data = [[Paragraph(num, toc_style), Paragraph(title, toc_style)] for num, title in toc_entries]
    toc_table = Table(toc_data, colWidths=[1*cm, 14.5*cm])
    toc_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
        ('FONTSIZE', (0, 0), (-1, -1), 14),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (0, -1), 0),
        ('LEFTPADDING', (1, 0), (1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ('LINEBELOW', (0, 0), (-1, -1), 0, black),
        ('LINEABOVE', (0, 0), (-1, -1), 0, black),
        ('LINEBEFORE', (0, 0), (-1, -1), 0, black),
        ('LINEAFTER', (0, 0), (-1, -1), 0, black),
    ]))
    elements.append(toc_table)
    elements.append(PageBreak())

    # ═══ KIRISH ═══
    elements.append(Paragraph("INTRODUCTION", chapter_style))
    elements.append(LINE)
    for para in work_data['introduction'].split('\n'):
        if para.strip():
            elements.append(Paragraph(para.strip(), base_style))

    # ═══ BOBLAR ═══
    diagrams = work_data.get('diagrams', [])
    extras = work_data.get('extras', [])
    chapter_diagrams = [
        [diagrams[0] if len(diagrams) > 0 else None, diagrams[1] if len(diagrams) > 1 else None],
        [diagrams[2] if len(diagrams) > 2 else None],
        [diagrams[3] if len(diagrams) > 3 else None],
    ]
    chapter_extras = [
        extras[0] if len(extras) > 0 else None,
        extras[1] if len(extras) > 1 else None,
        extras[2] if len(extras) > 2 else None,
    ]
    conclusion_diagram = diagrams[4] if len(diagrams) > 4 else None
    extra_for_conclusion = extras[3] if len(extras) > 3 else None

    fig_counter = [1]

    def add_image_with_caption(path, caption_text):
        if path and os.path.exists(path):
            try:
                elements.append(LINE)
                img = RLImage(path, width=15*cm, height=8.5*cm)
                img.hAlign = 'CENTER'
                elements.append(img)
                elements.append(Paragraph(f"Figure {fig_counter[0]} — {caption_text}", caption_style))
                elements.append(LINE)
                fig_counter[0] += 1
            except Exception as e:
                print(f"Rasm xatosi: {e}")

    for i, chapter in enumerate(work_data['chapters'], 1):
        elements.append(PageBreak())
        elements.append(Paragraph(f"CHAPTER {i}", chapter_style))
        elements.append(LINE)
        elements.append(Paragraph(chapter['title'].upper(),
                                   ParagraphStyle('ChT', parent=base_style, fontName='Times-Bold',
                                                  fontSize=14, alignment=TA_CENTER, spaceAfter=0)))
        elements.append(LINE)

        raw_paragraphs = [p for p in chapter['content'].split('\n') if p.strip()]
        paragraphs = []
        for p in raw_paragraphs:
            p_lower = p.lower().strip()
            if re.match(r'^chapter\s+\d+', p_lower):
                continue
            if p_lower == chapter['title'].lower().strip():
                continue
            paragraphs.append(p)
        section_img_inserted = False

        for j, para in enumerate(paragraphs):
            elements.append(Paragraph(para.strip(), base_style))
            if not section_img_inserted and j == 3 and chapter.get('image') and os.path.exists(chapter['image']):
                add_image_with_caption(chapter['image'], chapter['title'][:60])
                section_img_inserted = True
            if j == len(paragraphs) // 2:
                for diag_path in chapter_diagrams[i-1]:
                    add_image_with_caption(diag_path, f"Diagram — {chapter['title'][:50]}")

        if not section_img_inserted and chapter.get('image') and os.path.exists(chapter['image']):
            add_image_with_caption(chapter['image'], chapter['title'][:60])

        add_image_with_caption(chapter_extras[i-1], f"Illustration — {chapter['title'][:50]}")

    # ═══ XULOSA ═══
    elements.append(PageBreak())
    elements.append(Paragraph("CONCLUSION", chapter_style))
    elements.append(LINE)
    for para in work_data['conclusion'].split('\n'):
        if para.strip():
            elements.append(Paragraph(para.strip(), base_style))

    add_image_with_caption(conclusion_diagram, "Summary Diagram")
    add_image_with_caption(extra_for_conclusion, "Additional Illustration")
    elements.append(PageBreak())

    # ═══ ADABIYOTLAR ═══
    elements.append(Paragraph("REFERENCES", chapter_style))
    elements.append(LINE)
    for ref in work_data['references'].split('\n'):
        if ref.strip():
            elements.append(Paragraph(ref.strip(), ref_style))

    doc.build(elements)
    print(f"PDF yaratildi: {pdf_path}")


# ─────────────────────────────────────────────────────
# DOCX YARATISH
# ─────────────────────────────────────────────────────
def add_toc_entry(doc, number: int, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run_num = p.add_run(f"{number}. ")
    run_num.font.name = 'Times New Roman'
    run_num.font.size = Pt(14)
    run_text = p.add_run(title.upper())
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(14)


def add_plain_paragraph(doc, text: str, size=14, indent_cm=0):
    p = doc.add_paragraph(text.strip())
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
    return p


def add_docx_image(doc, path, caption_text, fig_num):
    if path and os.path.exists(path):
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_after = Pt(0)
            p_img.paragraph_format.space_before = Pt(0)
            p_img.add_run().add_picture(path, width=Inches(5.5))
            p_cap = doc.add_paragraph(f"Figure {fig_num} — {caption_text}")
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(0)
            p_cap.paragraph_format.space_before = Pt(0)
            for run in p_cap.runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
                run.italic = True
            return fig_num + 1
        except Exception as e:
            print(f"DOCX rasm xatosi: {e}")
    return fig_num


def create_docx(work_data: dict, meta: dict, docx_path: str):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    def add_centered(text, bold=False, size=14):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'

    def add_section_heading(text, size=14):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'

    # ═══ TITUL VARAQ ═══
    add_centered("MINISTRY OF HIGHER EDUCATION, SCIENCE AND INNOVATIONS", size=12)
    add_centered("OF THE REPUBLIC OF UZBEKISTAN", size=12)
    doc.add_paragraph()
    add_centered(meta.get("university", "UNIVERSITY").upper(), bold=True, size=14)
    add_centered(meta.get("faculty", "FACULTY").upper(), bold=True, size=12)
    add_centered(f"Department: {meta.get('department', 'DEPARTMENT').upper()}", size=12)
    doc.add_paragraph()
    add_centered("INDEPENDENT WORK", bold=True, size=16)
    add_centered(f"Subject: {meta.get('subject', 'SUBJECT').upper()}", size=12)
    add_centered(f"Topic: {meta.get('topic', 'TOPIC').upper()}", bold=True, size=14)
    doc.add_paragraph()

    if work_data.get('main_image') and os.path.exists(work_data['main_image']):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.add_run().add_picture(work_data['main_image'], width=Inches(5.5))
        except Exception as e:
            print(f"DOCX asosiy rasm: {e}")

    doc.add_paragraph()
    add_centered(f"Performed by: {meta.get('student', 'STUDENT').upper()}", size=12)
    add_centered(f"Group: {meta.get('group', 'GROUP')}, Course: {meta.get('course', '1')}", size=12)
    add_centered(f"Supervised by: {meta.get('teacher', 'TEACHER').upper()}", size=12)
    doc.add_paragraph()
    add_centered(f"{meta.get('city', 'CITY')} - {datetime.now().year}", size=12)
    doc.add_page_break()

    # ═══ MUNDARIJA ═══
    add_centered("TABLE OF CONTENTS", bold=True, size=16)
    toc_num = 1
    add_toc_entry(doc, toc_num, "INTRODUCTION"); toc_num += 1
    for chapter in work_data['chapters']:
        add_toc_entry(doc, toc_num, chapter['title']); toc_num += 1
    add_toc_entry(doc, toc_num, "CONCLUSION"); toc_num += 1
    add_toc_entry(doc, toc_num, "REFERENCES")
    doc.add_page_break()

    # ═══ KIRISH ═══
    add_section_heading("INTRODUCTION")
    for para in work_data['introduction'].split('\n'):
        if para.strip():
            add_plain_paragraph(doc, para.strip())

    # ═══ BOBLAR ═══
    diagrams = work_data.get('diagrams', [])
    extras = work_data.get('extras', [])
    chapter_diagrams = [
        [diagrams[0] if len(diagrams) > 0 else None, diagrams[1] if len(diagrams) > 1 else None],
        [diagrams[2] if len(diagrams) > 2 else None],
        [diagrams[3] if len(diagrams) > 3 else None],
    ]
    chapter_extras = [
        extras[0] if len(extras) > 0 else None,
        extras[1] if len(extras) > 1 else None,
        extras[2] if len(extras) > 2 else None,
    ]
    conclusion_diagram = diagrams[4] if len(diagrams) > 4 else None
    extra_for_conclusion = extras[3] if len(extras) > 3 else None

    fig_counter = [1]

    def docx_add_img(path, caption):
        new_num = add_docx_image(doc, path, caption, fig_counter[0])
        fig_counter[0] = new_num

    for i, chapter in enumerate(work_data['chapters'], 1):
        doc.add_page_break()
        add_section_heading(f"CHAPTER {i}")
        add_centered(chapter['title'].upper(), bold=True, size=14)

        raw_paragraphs = [p for p in chapter['content'].split('\n') if p.strip()]
        paragraphs = []
        for p in raw_paragraphs:
            p_lower = p.lower().strip()
            if re.match(r'^chapter\s+\d+', p_lower):
                continue
            if p_lower == chapter['title'].lower().strip():
                continue
            paragraphs.append(p)
        section_img_inserted = False

        for j, para in enumerate(paragraphs):
            add_plain_paragraph(doc, para.strip())
            if not section_img_inserted and j == 3 and chapter.get('image') and os.path.exists(chapter['image']):
                docx_add_img(chapter['image'], chapter['title'][:60])
                section_img_inserted = True
            if j == len(paragraphs) // 2:
                for diag_path in chapter_diagrams[i-1]:
                    docx_add_img(diag_path, f"Diagram — {chapter['title'][:50]}")

        if not section_img_inserted and chapter.get('image') and os.path.exists(chapter['image']):
            docx_add_img(chapter['image'], chapter['title'][:60])

        docx_add_img(chapter_extras[i-1], f"Illustration — {chapter['title'][:50]}")

    # ═══ XULOSA ═══
    doc.add_page_break()
    add_section_heading("CONCLUSION")
    for para in work_data['conclusion'].split('\n'):
        if para.strip():
            add_plain_paragraph(doc, para.strip())

    docx_add_img(conclusion_diagram, "Summary Diagram")
    docx_add_img(extra_for_conclusion, "Additional Illustration")
    doc.add_page_break()

    # ═══ ADABIYOTLAR ═══
    add_section_heading("REFERENCES")
    for ref in work_data['references'].split('\n'):
        if ref.strip():
            add_plain_paragraph(doc, ref.strip(), size=12, indent_cm=0.5)

    doc.save(docx_path)
    print(f"DOCX yaratildi: {docx_path}")


# ─────────────────────────────────────────────────────
# API ROUTES
# ─────────────────────────────────────────────────────
@app.get("/")
async def root():
    return {"message": "Mustaqil Ish Generator API is running", "status": "ok"}

@app.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.post("/api/generate")
async def generate_document(request: Request):
    try:
        data = await request.json()

        if not data.get("subject") or not data.get("topic"):
            return JSONResponse({"success": False, "error": "Subject va Topic majburiy!"}, status_code=400)

        if not GITHUB_TOKEN:
            return JSONResponse({"success": False, "error": "GITHUB_TOKEN sozlanmagan!"}, status_code=500)

        file_id = uuid.uuid4().hex
        pdf_path = f"{TEMP_DIR}/{file_id}.pdf"
        docx_path = f"{TEMP_DIR}/{file_id}.docx"

        print(f"GENERATSIYA: {data['topic']}")

        work_data = generate_full_independent_work(
            data["subject"],
            data["topic"],
            data.get("cefr", "B2")
        )

        total_images = (
            sum(1 for c in work_data['chapters'] if c.get('image') and os.path.exists(c['image'])) +
            sum(1 for d in work_data.get('diagrams', []) if d and os.path.exists(d)) +
            sum(1 for e in work_data.get('extras', []) if e and os.path.exists(e)) +
            (1 if work_data.get('main_image') and os.path.exists(work_data['main_image']) else 0)
        )

        print(f"Jami: {work_data['total_words']} so'z, {total_images} rasm")

        create_pdf(work_data, data, pdf_path)
        create_docx(work_data, data, docx_path)

        topic_safe = re.sub(r'[^\w\-_]', '_', data["topic"])[:30]

        return JSONResponse({
            "success": True,
            "total_words": work_data['total_words'],
            "total_images": total_images,
            "chapters": len(work_data['chapters']),
            "pdf": {"url": f"/download/{file_id}.pdf", "name": f"{topic_safe}.pdf"},
            "word": {"url": f"/download/{file_id}.docx", "name": f"{topic_safe}.docx"}
        })

    except Exception as e:
        print(f"XATOLIK: {traceback.format_exc()}")
        return JSONResponse({"success": False, "error": str(e)}, status_code=500)


@app.get("/download/{filename}")
async def download_file(filename: str):
    filepath = f"{TEMP_DIR}/{filename}"
    if os.path.exists(filepath):
        media_type = (
            'application/pdf' if filename.endswith('.pdf')
            else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        return FileResponse(filepath, filename=filename, media_type=media_type)
    raise HTTPException(status_code=404, detail="Fayl topilmadi yoki muddati tugagan")